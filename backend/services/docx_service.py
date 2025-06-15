# AudioBook Organizer - DOCX Processing Service

from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tempfile
import os
import re
from typing import Dict, List, Any, Tuple


class DocxService:
    """Service for processing DOCX files and extracting text with formatting"""
    
    def __init__(self):
        # Map DOCX styles to AudioBook CSS classes
        self.style_mapping = {
            'Heading 1': 'title',
            'Title': 'title',
            'Heading 2': 'subtitle', 
            'Subtitle': 'subtitle',
            'Heading 3': 'section',
            'Heading 4': 'subsection',
            'Quote': 'quote',
            'Block Text': 'quote',
            'Intense Quote': 'quote'
        }
        
        # Font size thresholds for dynamic heading detection
        self.size_thresholds = {
            18: 'title',     # 18pt+ = title
            16: 'subtitle',  # 16pt+ = subtitle  
            14: 'section',   # 14pt+ = section
            12: 'subsection' # 12pt+ = subsection
        }
    
    def extract_content_with_formatting(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text and formatting from DOCX file with enhanced whitespace preservation
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dict containing text, formatting_ranges, and comments
        """
        try:
            doc = Document(file_path)
            
            result = {
                'text': '',
                'formatting_ranges': [],
                'comments': [],
                'metadata': {
                    'total_paragraphs': len(doc.paragraphs),
                    'processing_notes': []
                }
            }
            
            # Enhanced text extraction with better whitespace preservation
            extracted_data = self._extract_text_with_structure(doc)
            result['text'] = extracted_data['text']
            result['formatting_ranges'] = extracted_data['formatting_ranges']
            
            # Add metadata
            result['metadata']['total_formatting_ranges'] = len(result['formatting_ranges'])
            result['metadata']['final_text_length'] = len(result['text'])
            result['metadata']['processing_notes'] = extracted_data['processing_notes']
            
            return result
            
        except Exception as e:
            raise Exception(f"Failed to process DOCX file: {str(e)}")
    
    def _extract_text_with_structure(self, doc) -> Dict[str, Any]:
        """
        Enhanced text extraction that preserves DOCX structure and whitespace
        """
        text_parts = []
        formatting_ranges = []
        processing_notes = []
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            # Get paragraph start position in final text
            para_start_pos = len(''.join(text_parts))
            
            # Handle empty paragraphs - preserve them as line breaks
            if not paragraph.text.strip():
                # Add single newline for empty paragraph
                text_parts.append('\n')
                continue
            
            # Extract runs with enhanced formatting detection
            para_text, para_formatting = self._process_paragraph_enhanced(
                paragraph, para_start_pos, processing_notes
            )
            
            # Add paragraph text
            text_parts.append(para_text)
            
            # Add paragraph-level formatting
            formatting_ranges.extend(para_formatting)
            
            # Add newline after paragraph (except for last paragraph)
            if para_idx < len(doc.paragraphs) - 1:
                text_parts.append('\n')
        
        # Join all text parts
        final_text = ''.join(text_parts)
        
        # Validate and adjust formatting ranges
        validated_ranges = self._validate_formatting_ranges(
            formatting_ranges, final_text, processing_notes
        )
        
        return {
            'text': final_text,
            'formatting_ranges': validated_ranges,
            'processing_notes': processing_notes
        }
    
    def _process_paragraph_enhanced(self, paragraph, para_start_pos: int, processing_notes: List[str]) -> Tuple[str, List[Dict]]:
        """
        Enhanced paragraph processing with better run handling
        """
        para_text_parts = []
        para_formatting = []
        current_pos = para_start_pos
        
        # Process each run in the paragraph
        for run_idx, run in enumerate(paragraph.runs):
            if not run.text:
                continue
            
            # Get run text with preserved whitespace
            run_text = self._preserve_run_whitespace(run.text)
            run_start = current_pos
            run_end = current_pos + len(run_text)
            
            # Extract character-level formatting
            run_formatting = self._extract_run_formatting(
                run, run_start, run_end, processing_notes
            )
            para_formatting.extend(run_formatting)
            
            # Add run text
            para_text_parts.append(run_text)
            current_pos = run_end
        
        # Add paragraph-level formatting if applicable
        para_text = ''.join(para_text_parts)
        para_end = para_start_pos + len(para_text)
        
        paragraph_formatting = self._extract_paragraph_formatting(
            paragraph, para_start_pos, para_end, processing_notes
        )
        para_formatting.extend(paragraph_formatting)
        
        return para_text, para_formatting
    
    def _preserve_run_whitespace(self, text: str) -> str:
        """
        Preserve whitespace in runs while normalizing line endings
        """
        # Only normalize line endings, preserve all other whitespace
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Convert multiple consecutive spaces to single spaces (Word behavior)
        # but preserve intentional spacing
        text = re.sub(r'[ \t]+', ' ', text)
        
        return text
    
    def _extract_run_formatting(self, run, start: int, end: int, processing_notes: List[str]) -> List[Dict]:
        """Enhanced run formatting extraction with better detection"""
        formatting_ranges = []
        
        if start >= end:
            return formatting_ranges
        
        # Bold formatting
        if run.bold:
            formatting_ranges.append({
                'start': start,
                'end': end,
                'type': 'bold',
                'level': 1,
                'source': 'run_bold'
            })
        
        # Italic formatting  
        if run.italic:
            formatting_ranges.append({
                'start': start,
                'end': end,
                'type': 'italic',
                'level': 1,
                'source': 'run_italic'
            })
        
        # Underline formatting
        if run.underline:
            formatting_ranges.append({
                'start': start,
                'end': end,
                'type': 'underline',
                'level': 1,
                'source': 'run_underline'
            })
        
        # Enhanced font size-based heading detection
        if hasattr(run.font, 'size') and run.font.size:
            size_pt = run.font.size.pt
            heading_type = self._determine_heading_from_size(size_pt)
            if heading_type:
                formatting_ranges.append({
                    'start': start,
                    'end': end,
                    'type': heading_type,
                    'level': 1,
                    'source': f'font_size_{size_pt}pt'
                })
                processing_notes.append(f'Applied {heading_type} from font size: {size_pt}pt')
        
        return formatting_ranges
    
    def _extract_paragraph_formatting(self, paragraph, start: int, end: int, processing_notes: List[str]) -> List[Dict]:
        """Enhanced paragraph formatting extraction"""
        formatting_ranges = []
        
        if start >= end:
            return formatting_ranges
        
        # Style-based formatting with better mapping
        style_name = paragraph.style.name
        if style_name in self.style_mapping:
            format_type = self.style_mapping[style_name]
            formatting_ranges.append({
                'start': start,
                'end': end,
                'type': format_type,
                'level': 1,
                'source': f'paragraph_style_{style_name}'
            })
            processing_notes.append(f'Applied {format_type} from style: {style_name}')
        
        # Enhanced alignment-based formatting
        if hasattr(paragraph, 'alignment') and paragraph.alignment:
            if paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                formatting_ranges.append({
                    'start': start,
                    'end': end,
                    'type': 'quote',
                    'level': 1,
                    'source': 'center_alignment'
                })
            elif paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                formatting_ranges.append({
                    'start': start,
                    'end': end,
                    'type': 'quote',
                    'level': 1,
                    'source': 'right_alignment'
                })
        
        return formatting_ranges
    
    def _determine_heading_from_size(self, size_pt: float) -> str:
        """Determine heading type from font size"""
        for threshold, heading_type in sorted(self.size_thresholds.items(), reverse=True):
            if size_pt >= threshold:
                return heading_type
        return None
    
    def _validate_formatting_ranges(self, formatting_ranges: List[Dict], text: str, processing_notes: List[str]) -> List[Dict]:
        """
        Validate and fix formatting ranges against final text
        """
        text_length = len(text)
        valid_ranges = []
        
        processing_notes.append(f"Validating {len(formatting_ranges)} ranges against text length {text_length}")
        
        for range_obj in formatting_ranges:
            start = range_obj['start']
            end = range_obj['end']
            
            # Ensure valid bounds
            if start < 0:
                start = 0
            if end > text_length:
                end = text_length
            if start >= end:
                processing_notes.append(f"Skipped invalid range {range_obj['start']}-{range_obj['end']}")
                continue
            
            # Create validated range
            validated_range = range_obj.copy()
            validated_range['start'] = start
            validated_range['end'] = end
            
            if start != range_obj['start'] or end != range_obj['end']:
                validated_range['source'] = f"{range_obj.get('source', 'unknown')}_bounds_adjusted"
                processing_notes.append(f"Adjusted range {range_obj['start']}-{range_obj['end']} -> {start}-{end}")
            
            valid_ranges.append(validated_range)
        
        # Sort ranges by position for consistent application
        valid_ranges.sort(key=lambda x: (x['start'], x['end'] - x['start']))
        
        processing_notes.append(f"Validation complete: {len(valid_ranges)}/{len(formatting_ranges)} ranges valid")
        return valid_ranges
    
    def validate_docx_file(self, file_path: str) -> Dict[str, Any]:
        """
        Validate DOCX file before processing
        
        Returns:
            Dict with validation results
        """
        try:
            # Try to open the document
            doc = Document(file_path)
            
            return {
                'valid': True,
                'paragraphs': len(doc.paragraphs),
                'has_content': any(p.text.strip() for p in doc.paragraphs),
                'styles_found': list(set(p.style.name for p in doc.paragraphs)),
                'estimated_size': 'small' if len(doc.paragraphs) < 100 else 'medium' if len(doc.paragraphs) < 500 else 'large'
            }
            
        except Exception as e:
            return {
                'valid': False,
                'error': str(e),
                'error_type': type(e).__name__
            }
    
    def get_processing_info(self, file_path: str) -> Dict[str, Any]:
        """Get information about DOCX file for processing estimates"""
        try:
            validation = self.validate_docx_file(file_path)
            if not validation['valid']:
                return validation
            
            doc = Document(file_path)
            
            # Count formatting elements
            total_runs = sum(len(p.runs) for p in doc.paragraphs)
            formatted_runs = sum(
                1 for p in doc.paragraphs 
                for r in p.runs 
                if r.bold or r.italic or r.underline or (hasattr(r.font, 'size') and r.font.size)
            )
            
            return {
                'valid': True,
                'paragraphs': len(doc.paragraphs),
                'total_runs': total_runs,
                'formatted_runs': formatted_runs,
                'formatting_density': formatted_runs / max(total_runs, 1),
                'estimated_processing_time': self._estimate_processing_time(len(doc.paragraphs), formatted_runs),
                'styles_used': validation['styles_found']
            }
            
        except Exception as e:
            return {
                'valid': False,
                'error': str(e)
            }
    
    def _estimate_processing_time(self, paragraphs: int, formatted_runs: int) -> str:
        """Estimate processing time based on document complexity"""
        complexity_score = paragraphs * 0.1 + formatted_runs * 0.5
        
        if complexity_score < 10:
            return '< 1 second'
        elif complexity_score < 50:
            return '1-3 seconds'
        elif complexity_score < 200:
            return '3-10 seconds'
        else:
            return '10+ seconds'

    def extract_text_only(self, file_path: str) -> Dict[str, Any]:
        """
        Extract ONLY plain text from DOCX file, preserving whitespace and line breaks
        NO formatting extraction - just clean text
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dict containing only text and basic metadata
        """
        try:
            doc = Document(file_path)
            
            text_parts = []
            
            for para_idx, paragraph in enumerate(doc.paragraphs):
                # Get paragraph text exactly as it is
                para_text = paragraph.text
                
                # Add paragraph text (even if empty)
                text_parts.append(para_text)
                
                # Add newline after paragraph (except for last paragraph)
                if para_idx < len(doc.paragraphs) - 1:
                    text_parts.append('\n')
            
            # Join all text parts
            final_text = ''.join(text_parts)
            
            return {
                'text': final_text,
                'formatting_ranges': [],  # No formatting - empty array
                'comments': [],
                'metadata': {
                    'total_paragraphs': len(doc.paragraphs),
                    'final_text_length': len(final_text),
                    'processing_method': 'text_only',
                    'processing_notes': ['Extracted text only, no formatting applied']
                }
            }
            
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX file: {str(e)}") 